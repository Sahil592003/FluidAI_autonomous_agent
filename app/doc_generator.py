
"""Universal document generator for ANY document type."""

import os
import logging
from datetime import datetime
from typing import List, Optional
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.config import Config
from app.models import ExecutionContext
from app.prompts.prompts import UNIVERSAL_SECTION_TEMPLATES

logger = logging.getLogger(__name__)


class DocumentGenerator:
    """Generate professional documents for ANY type."""
    
    def __init__(self):
        Config.ensure_docs_dir()
        self.templates = UNIVERSAL_SECTION_TEMPLATES
    
    def generate(self, context: ExecutionContext) -> str:
        """Generate a document for ANY type."""
        logger.info(f"Generating {context.document_type}")
        
        try:
            doc = Document()
            
            # Add title with document type
            self._add_title(doc, context)
            
            # Add metadata
            self._add_metadata(doc, context)
            
            # Add content sections
            self._add_sections(doc, context)
            
            # Add professional footer
            self._add_footer(doc, context)
            
            # Save document
            filename = self._get_filename(context)
            filepath = os.path.join(Config.DOCS_OUTPUT_DIR, filename)
            doc.save(filepath)
            
            logger.info(f"✅ Document saved: {filepath}")
            return filepath
            
        except Exception as e:
            logger.error(f"Document generation failed: {str(e)}")
            raise
    
    def _add_title(self, doc: Document, context: ExecutionContext) -> None:
        """Add professional title."""
        # Main title
        title = f"{context.document_type or 'Document'}"
        title_para = doc.add_heading(title, 0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Subtitle
        subtitle = doc.add_paragraph(f"Generated: {datetime.now().strftime('%B %d, %Y')}")
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle.runs[0].font.size = Pt(12)
        
        # Divider
        doc.add_paragraph('_' * 80)
        doc.add_paragraph()
    
    def _add_metadata(self, doc: Document, context: ExecutionContext) -> None:
        """Add document metadata."""
        meta_para = doc.add_paragraph()
        meta_para.add_run(f"Document Type: {context.document_type or 'Document'}\n").bold = True
        meta_para.add_run(f"Created: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        
        # Add assumptions if any
        if context.assumptions:
            doc.add_paragraph("\nAssumptions:")
            for assumption in context.assumptions[:3]:
                doc.add_paragraph(f"• {assumption}", style='List Bullet')
        
        doc.add_paragraph()
        doc.add_paragraph('_' * 80)
        doc.add_paragraph()
    
    def _add_sections(self, doc: Document, context: ExecutionContext) -> None:
        """Add all document sections."""
        for i, result in enumerate(context.results, 1):
            if result.status == "completed" and result.content:
                # Extract section name
                section_name = result.task.split(":", 1)[-1].strip() if ":" in result.task else f"Section {i}"
                
                # Add heading
                doc.add_heading(f"{i}. {section_name}", 1)
                
                # Add content
                paragraphs = result.content.split('\n')
                for para in paragraphs:
                    if para.strip():
                        p = doc.add_paragraph()
                        p.add_run(para.strip()).font.size = Pt(11)
                
                doc.add_paragraph()
    
    def _add_footer(self, doc: Document, context: ExecutionContext) -> None:
        """Add professional footer."""
        for section in doc.sections:
            footer = section.footer
            footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            footer_text = f"{context.document_type or 'Document'} | Page "
            footer_para.add_run(footer_text).font.size = Pt(9)
            
            # Add page number field
            footer_para.add_run("{PAGE}").font.size = Pt(9)
    
    def _get_filename(self, context: ExecutionContext) -> str:
        """Generate filename for ANY document."""
        doc_type = context.document_type or "document"
        clean_type = doc_type.lower().replace(' ', '_').replace('/', '_')
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        return f"{clean_type}_{timestamp}.docx"